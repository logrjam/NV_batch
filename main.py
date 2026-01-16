# -*- coding: utf-8 -*-
"""
This is the main script which runs from the batch file "GenerateReport"
requirements.txt must exist with the correct packages 
config.txt must have values for ReportYear and ReportMonth
data_processing folder must exist with internal modules


Author: Logan Jamison
Date: 1/15/2026
"""
import os

def load_config(path="config.txt"):
    values = {}
    with open(path) as f:
        for line in f:
            if "=" in line:
                key, value = line.strip().split("=", 1)
                values[key] = value
    return values

def main():
    import pandas as pd
    import matplotlib.pyplot as plt
    from matplotlib.dates import DateFormatter
    import numpy as np
    from data_processing.LoadBasicReportData_WSOR_NV import LoadBasicReportData_WSOR_NV
    from data_processing.LoadReportData_NV import LoadReportData_NV
    import datetime as dt
    from docx import Document
    from docx2pdf import convert
    import matplotlib.gridspec as gridspec
    from docx.shared import Cm
    from docx.shared import Pt
    from data_processing.add_hyperlink import add_hyperlink
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    import math

    
    
    

    cfg = load_config()
    ReportYear = int(cfg["ReportYear"])
    ReportMonth = int(cfg["ReportMonth"])

    print(f"Generating report for {ReportYear}-{ReportMonth}")
    # your pipeline logic here


            
    watershedlist = [  
                            'state_of_nevada_and_eastern_sierra',
                            'lake_tahoe',
                            'truckee',
                            'carson',
                            'walker',
                            'northern_great_basin',
                            'upper_humboldt',
                            'lower_humboldt',
                            'clover_valley_and_franklin',
                            'snake',
                            'owyhee',
                            'eastern_nevada',
                            'spring_mountains',
                            'surprise_valley-warner_mtns',
                            'Upper_Colorado_Region'
                      ]
    
    
    
    [prec_df_api,basin_res_df_api,res_d] = LoadBasicReportData_WSOR_NV(ReportMonth,ReportYear)
    
    fname_append = []
    
    for i in range(len(watershedlist)):
        
        wshed = watershedlist[i]
            
            
        [pcp,swe,moi,res,pcp_monthly,txt,fname1,cleantxt,curswepct,curpcppct,smcur]=LoadReportData_NV(wshed,ReportYear,ReportMonth,prec_df_api,basin_res_df_api,res_d)
        print("Generating Summary Plot for " + cleantxt)
     
        fname_append.append(fname1)                
        
        curpcp = pcp['2019-10-01':'2020'+'-0'+str(ReportMonth)+'-01']
        curswe = swe['2019-10-01':'2020'+'-0'+str(ReportMonth)+'-01']
        
        if len(moi)!=0:
            curmoi = moi['2019-10-01':'2020'+'-0'+str(ReportMonth)+'-01']
        
        # curpcp = pcp[str(wy-2)+'-10-01':str(wy-1)+'-0'+str(ReportMonth)+'-01']
        # curswe = swe[str(wy-2)+'-10-01':str(wy-1)+'-0'+str(ReportMonth)+'-01']
        # curmoi = moi[str(wy-2)+'-10-01':str(wy-1)+'-0'+str(ReportMonth)+'-01']
        
        if len(res) != 0:
            if wshed!='state_of_nevada_and_eastern_sierra':
             
                res = res.sort_values('res_cap', ascending = False)
                
                reslabels = res.index.tolist()
                reslabels = [w.replace('Reservoir', '') for w in reslabels]
                
                # make replacement text for res labels as needed
                if wshed == 'lake_tahoe':
                    replacements = {'Marlette Lk': 'Marlette Lake'}
                elif wshed == 'truckee':
                    replacements = {'Prosser': 'Prosser\nReservoir',
                                    'Independence': 'Independence\nLake',
                                    'Stampede': 'Stampede\nReservoir',
                                    'Boca': 'Boca\nReservoir'}
                elif wshed == 'walker':
                    replacements = {'Topaz Lk': 'Topaz Lake',
                                    'Bridgeport': 'Bridgeport\nReservoir'}
                elif wshed == 'lower_humboldt':
                    replacements = {'Rye Patch Re':'Rye Patch\nReservoir',
                                    'Chimney Creek': 'Chimney Ck.\nReservoir'}
                elif wshed == 'carson':
                    replacements = {'Lahontan': 'Lahontan\nReservoir'}
                elif wshed == 'owyhee':
                    replacements ={'Wild Horse': 'Wildhorse\nReservoir'} 
                else:
                    replacements ={} 
                    
                # make the replacements for reslabels as needed    
                if len(replacements) != 0:
                    reslabels = [
                        next((new for old, new in replacements.items() if old in item), item)
                        for item in reslabels
                    ]
                    
            else:
    
                lake_tahoe = ['Lake Tahoe']
                truckee_res = ['Donner Lake','Prosser Reservoir','Independence Lake','Stampede Reservoir','Boca Reservoir']    
                carson_res = ['Lahontan Reservoir']
                walker_res = ['Topaz Lk nr Topaz','Bridgeport Reservoir']
                humboldt_res = ['Rye Patch Re nr Rye Patch, NV','Chimney Creek Reservoir']
                owyhee_res = ['Wild Horse Reservoir']
                lowercolorado_res = ['Lake Mohave','Lake Mead','Lake Powell']
                
                mydict = {
                    'Watershed/Region': 
                         ['Lake Tahoe','Truckee\nBasin','Carson Basin','Walker Basin','Humboldt\nBasin','Owyhee\nBasin','Lwr. Colorado\nBasin'],  
                    
                "storage_basin_prev": 
                    [
                      int(res.loc[lake_tahoe,'res_ly'].sum()/res.loc[lake_tahoe,'res_cap'].sum()*100),\
                      int(res.loc[truckee_res,'res_ly'].sum()/res.loc[truckee_res,'res_cap'].sum()*100),\
                      int(res.loc[carson_res,'res_ly'].sum()/res.loc[carson_res,'res_cap'].sum()*100),\
                      int(res.loc[walker_res,'res_ly'].sum()/res.loc[walker_res,'res_cap'].sum()*100),\
                      int(res.loc[humboldt_res,'res_ly'].sum()/res.loc[humboldt_res,'res_cap'].sum()*100),\
                      int(res.loc[owyhee_res,'res_ly'].sum()/res.loc[owyhee_res,'res_cap'].sum()*100),\
                      int(res.loc[lowercolorado_res,'res_ly'].sum()/res.loc[lowercolorado_res,'res_cap'].sum()*100)],
                   
                "storage_basin_cur": 
                    [
                      int(res.loc[lake_tahoe,'res_curr'].sum()/res.loc[lake_tahoe,'res_cap'].sum()*100),\
                      int(res.loc[truckee_res,'res_curr'].sum()/res.loc[truckee_res,'res_cap'].sum()*100),\
                      int(res.loc[carson_res,'res_curr'].sum()/res.loc[carson_res,'res_cap'].sum()*100),\
                      int(res.loc[walker_res,'res_curr'].sum()/res.loc[walker_res,'res_cap'].sum()*100),\
                      int(res.loc[humboldt_res,'res_curr'].sum()/res.loc[humboldt_res,'res_cap'].sum()*100),\
                      int(res.loc[owyhee_res,'res_curr'].sum()/res.loc[owyhee_res,'res_cap'].sum()*100),\
                      int(res.loc[lowercolorado_res,'res_curr'].sum()/res.loc[lowercolorado_res,'res_cap'].sum()*100)]
                        }
                          
                res_basin_df = pd.DataFrame (mydict)
                res_basin_df.index = res_basin_df['Watershed/Region']
                res_basin_df.drop(res_basin_df.columns[0],axis=1,inplace=True)
                reslabels = res_basin_df.index.tolist()
                
                # res_basin_df = pd.DataFrame (mydict, columns = ['storage_basin_prev','storage_basin_cur'], index = ['Lake Tahoe','Truckee\nBasin','Carson Basin','Walker Basin','Humboldt\nBasin','Owyhee\nBasin','Lwr. Colorado\nBasin'])
                # reslabels = res_basin_df.index.tolist()
       
        
        date = pcp.index
    
        #If there are reservoirs do this!
        if len(res) !=0 or wshed=='clover_valley_and_franklin':
            fig = plt.figure(1)
            fig.clf()
            fig.set_size_inches(9,9)
            ax2 = fig.add_subplot(223)
            ax2.fill_between(pcp.index,pcp.Min,pcp['10%'],facecolor=(253/255, 217/255, 217/255))
            ax2.fill_between(pcp.index,pcp['10%'],pcp['30%'],facecolor=(253/255, 253/255, 217/255))
            ax2.fill_between(pcp.index,pcp['30%'],pcp['70%'],facecolor=(234/255, 253/255, 234/255))
            ax2.fill_between(pcp.index,pcp['70%'],pcp['90%'],facecolor=(217/255, 253/255, 253/255))
            ax2.fill_between(pcp.index,pcp['90%'],pcp['Max'],facecolor=(217/255, 217/255, 253/255))
            ax2.plot(pcp.index,pcp['Max'],color=(147/255, 147/255, 248/255),linewidth=1) #max
            ax2.plot(pcp.index,pcp['Min'],color=(245/255, 108/255, 108/255),linewidth=1) #min
            ax2.plot(pcp.index,pcp["Median ('91-'20)"],color=(199/255, 250/255, 199/255),linewidth=2,label='Current')
            ax2.plot(curpcp[pcp.columns[1]],color='black',linewidth=2,label='Current') #current
            ax2.set(ylabel = "Precipitation (Inches)")
            date_form = DateFormatter("%b")
            ax2.xaxis.set_major_formatter(date_form)
           
            
            
            ax2.text(0,1.01,'Precipitation',transform = ax2.transAxes,fontsize = 9,weight='bold') 
            ax2.text(.01,.95,'Water Year % Median: '+str(int(curpcppct)),transform = ax2.transAxes,fontsize = 9) 
            ax2.text(.01,.88,'Monthly % Median',transform = ax2.transAxes,fontsize = 9)       
            ax2.text(.01,.877,'___________________',transform = ax2.transAxes,fontsize = 9)      
            
            if ReportMonth == 1:
                ax2.text(.01,.75,\
                          '\nOct: {}'\
                          '\nNov: {}'\
                          '\nDec: {}'\
                          .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp))\
                          ,transform = ax2.transAxes\
                          ,fontsize = 8)                             
            elif ReportMonth == 2:
                ax2.text(.01,.71,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)    
            elif ReportMonth == 3:
                ax2.text(.01,.67,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                           
            elif ReportMonth == 4:
                ax2.text(.01,.63,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                            
            elif ReportMonth == 5:
                ax2.text(.01,.59,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            '\nApr: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp),int(pcp_monthly.Apr_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8) 
       
            
        
        elif wshed=='surprise_valley-warner_mtns':
            fig = plt.figure(1)
            fig.clf()
            fig.set_size_inches(9,9)
            ax2 = fig.add_subplot(222)
            ax2.fill_between(pcp.index,pcp.Min,pcp['10%'],facecolor=(253/255, 217/255, 217/255))
            ax2.fill_between(pcp.index,pcp['10%'],pcp['30%'],facecolor=(253/255, 253/255, 217/255))
            ax2.fill_between(pcp.index,pcp['30%'],pcp['70%'],facecolor=(234/255, 253/255, 234/255))
            ax2.fill_between(pcp.index,pcp['70%'],pcp['90%'],facecolor=(217/255, 253/255, 253/255))
            ax2.fill_between(pcp.index,pcp['90%'],pcp['Max'],facecolor=(217/255, 217/255, 253/255))
            ax2.plot(pcp.index,pcp['Max'],color=(147/255, 147/255, 248/255),linewidth=1) #max
            ax2.plot(pcp.index,pcp['Min'],color=(245/255, 108/255, 108/255),linewidth=1) #min
            ax2.plot(pcp.index,pcp["Median ('91-'20)"],color=(199/255, 250/255, 199/255),linewidth=2,label='Current')
            ax2.plot(curpcp[pcp.columns[1]],color='black',linewidth=2,label='Current') #current
            ax2.set(ylabel = "Precipitation (Inches)")
            date_form = DateFormatter("%b")
            ax2.xaxis.set_major_formatter(date_form)
             
            
            ax2.text(0,1.01,'Precipitation',transform = ax2.transAxes,fontsize = 9,weight='bold') 
            ax2.text(.01,.95,'Water Year % Median: '+str(int(curpcppct)),transform = ax2.transAxes,fontsize = 9) 
            ax2.text(.01,.88,'Monthly % Median',transform = ax2.transAxes,fontsize = 9)       
            ax2.text(.01,.877,'___________________',transform = ax2.transAxes,fontsize = 9)          
              
        
            
            if ReportMonth == 1:
                ax2.text(.01,.75,\
                          '\nOct: {}'\
                          '\nNov: {}'\
                          '\nDec: {}'\
                          .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp))\
                          ,transform = ax2.transAxes\
                          ,fontsize = 8)                             
            elif ReportMonth == 2:
                ax2.text(.01,.71,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)    
            elif ReportMonth == 3:
                ax2.text(.01,.67,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                           
            elif ReportMonth == 4:
                ax2.text(.01,.63,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                            
            elif ReportMonth == 5:
                ax2.text(.01,.59,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            '\nApr: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp),int(pcp_monthly.Apr_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8) 
            
                
        else:
            gs = gridspec.GridSpec(4, 4)      
            fig = plt.figure(1)
            fig.clf()
            fig.set_size_inches(9,9)
            ax2 = fig.add_subplot(gs[2:4, 1:3])
            ax2.fill_between(pcp.index,pcp.Min,pcp['10%'],facecolor=(253/255, 217/255, 217/255))
            ax2.fill_between(pcp.index,pcp['10%'],pcp['30%'],facecolor=(253/255, 253/255, 217/255))
            ax2.fill_between(pcp.index,pcp['30%'],pcp['70%'],facecolor=(234/255, 253/255, 234/255))
            ax2.fill_between(pcp.index,pcp['70%'],pcp['90%'],facecolor=(217/255, 253/255, 253/255))
            ax2.fill_between(pcp.index,pcp['90%'],pcp['Max'],facecolor=(217/255, 217/255, 253/255))
            ax2.plot(pcp.index,pcp['Max'],color=(147/255, 147/255, 248/255),linewidth=1) #max
            ax2.plot(pcp.index,pcp['Min'],color=(245/255, 108/255, 108/255),linewidth=1) #min
            ax2.plot(pcp.index,pcp["Median ('91-'20)"],color=(199/255, 250/255, 199/255),linewidth=2,label='Current')
            ax2.plot(curpcp[pcp.columns[1]],color='black',linewidth=2,label='Current') #current
            ax2.set(ylabel = "Precipitation (Inches)")
            date_form = DateFormatter("%b")
            ax2.xaxis.set_major_formatter(date_form)
             
            
            ax2.text(0,1.01,'Precipitation',transform = ax2.transAxes,fontsize = 9,weight='bold') 
            ax2.text(.01,.95,'Water Year % Median: '+str(int(curpcppct)),transform = ax2.transAxes,fontsize = 9) 
            ax2.text(.01,.88,'Monthly % Median',transform = ax2.transAxes,fontsize = 9)       
            ax2.text(.01,.877,'___________________',transform = ax2.transAxes,fontsize = 9)          
    
            if ReportMonth == 1:
                ax2.text(.01,.75,\
                          '\nOct: {}'\
                          '\nNov: {}'\
                          '\nDec: {}'\
                          .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp))\
                          ,transform = ax2.transAxes\
                          ,fontsize = 8)                             
            elif ReportMonth == 2:
                ax2.text(.01,.71,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)    
            elif ReportMonth == 3:
                ax2.text(.01,.67,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                           
            elif ReportMonth == 4:
                ax2.text(.01,.63,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8)                            
            elif ReportMonth == 5:
                ax2.text(.01,.59,\
                            '\nOct: {}'\
                            '\nNov: {}'\
                            '\nDec: {}'\
                            '\nJan: {}'\
                            '\nFeb: {}'\
                            '\nMar: {}'\
                            '\nApr: {}'\
                            .format(int(pcp_monthly.Oct_pcp),int(pcp_monthly.Nov_pcp),int(pcp_monthly.Dec_pcp),int(pcp_monthly.Jan_pcp),int(pcp_monthly.Feb_pcp),int(pcp_monthly.March_pcp),int(pcp_monthly.Apr_pcp))\
                            ,transform = ax2.transAxes\
                            ,fontsize = 8) 
        
        
        #Clover valley note
        if wshed=='clover_valley_and_franklin':
            ax2.text(1.1,.95, 'Hole-in-Mountain SNOTEL - Status',transform = ax2.transAxes,fontsize = 9,weight='bold')
            ax2.text(1.1,.1,'Starting in water year 2020, automated snow\n'\
                            'water and snow depth measurements have been\n'\
                            'moved back to the original Hole-in-Mountain SNOTEL\n'\
                            'location used from 1981-2015. This move allows \n'\
                            'daily snow water percent of median to be calculated\n'\
                            'using historic data. The SNOTEL was re-located\n'\
                            'outside an avalanche zone in 2016 following an\n'\
                            'avalanche that damaged the site. Unfortunately, the\n'\
                            'new location while protected from future slides, was\n'\
                            'subject to drifting and snow data proved\n'\
                            'unrepresentative. Snow data from 2016-2020 have\n'\
                            'been removed from the public database and will\n'\
                            'appear as missing in NRCS products. Other SNOTEL\n'\
                            'parameters collected at the newer location are\n'\
                            'representative and were not moved. These include\n'\
                            'air temperature, precipitation and soil moisture.\n\n'\
                            'Contact Jeff Anderson for more information:\n'\
                            'jeff.anderson@usda.gov or 775-834-0913',transform = ax2.transAxes,fontsize = 9)
                     
        
                
        ax2 = fig.add_subplot(221)   
        ax2.fill_between(swe.index,swe.Min,swe['10%'],facecolor=(253/255, 217/255, 217/255))
        ax2.fill_between(swe.index,swe['10%'],swe['30%'],facecolor=(253/255, 253/255, 217/255))
        ax2.fill_between(swe.index,swe['30%'],swe['70%'],facecolor=(234/255, 253/255, 234/255))
        ax2.fill_between(swe.index,swe['70%'],swe['90%'],facecolor=(217/255, 253/255, 253/255))
        ax2.fill_between(swe.index,swe['90%'],swe['Max'],facecolor=(217/255, 217/255, 253/255))
        ax2.plot(swe.index,swe['Max'],color=(147/255, 147/255, 248/255),linewidth=1) #max
        ax2.plot(swe.index,swe['Min'],color=(245/255, 108/255, 108/255),linewidth=1) #min
        ax2.plot(swe.index,swe["Median ('91-'20)"],color=(199/255, 250/255, 199/255),linewidth=2,label='30 Yr\nMedian')
        ax2.plot(curswe[swe.columns[1]],color='black',linewidth=2,label='Current\nConditions') #current
        ax2.set(ylabel = "Snow Water Equivalent (Inches)")
        
        if (wshed== 'clover_valley_and_franklin' or wshed == 'spring_mtns' or wshed=='state_of_nevada_and_eastern_sierra' or wshed=='surprise_valley-warner_mtns'):
            ax2.legend(fontsize=8,loc='upper right', bbox_to_anchor=(1, 0.95),frameon=True)
        else:
            ax2.legend(fontsize=8,loc='upper right',frameon=True)
    
        
        curax = ax2.get_ylim()
        ax2.set_ylim(curax[0],curax[1]*1.03)
        
        
        date_form = DateFormatter("%b")
        ax2.xaxis.set_major_formatter(date_form)
    
        ax2.text(0,1.01,'Snowpack',transform = ax2.transAxes,fontsize = 9,weight='bold') 
        ax2.text(.01,.95,cleantxt,transform=ax2.transAxes,fontsize = 9)
        
        if curswepct!=-9999 and math.isnan(curswepct)!=True:   
            ax2.text(.01,.90,'% Median: '+ str(int(curswepct)),transform=ax2.transAxes,fontsize = 9)
        
        # if wshed!='state_of_nevada':
        #     ax2.text(0,1.01,'Snowpack',transform = ax2.transAxes,fontsize = 9,weight='bold') 
        #     ax2.text(.01,.83,'Subbasin % Normal',transform = ax2.transAxes,fontsize = 9)       
        #     ax2.text(.01,.827,'____________________',transform = ax2.transAxes,fontsize = 9)       
                   
        if len(moi)!=0:        
            ax2 = fig.add_subplot(222)
            ax2.fill_between(moi.index,moi.Min,moi['10%'],facecolor=(253/255, 217/255, 217/255))
            ax2.fill_between(moi.index,moi['10%'],moi['30%'],facecolor=(253/255, 253/255, 217/255))
            ax2.fill_between(moi.index,moi['30%'],moi['70%'],facecolor=(234/255, 253/255, 234/255))
            ax2.fill_between(moi.index,moi['70%'],moi['90%'],facecolor=(217/255, 253/255, 253/255))
            ax2.fill_between(moi.index,moi['90%'],moi['Max'],facecolor=(217/255, 217/255, 253/255))
            ax2.plot(moi.index,moi['Max'],color=(147/255, 147/255, 248/255),linewidth=1) #max
            ax2.plot(moi.index,moi['Min'],color=(245/255, 108/255, 108/255),linewidth=1) #min
            ax2.plot(moi.index,moi['Median (POR)'],color=(199/255, 250/255, 199/255),linewidth=2,label='Current')
            ax2.plot(curmoi[moi.columns[1]],color='black',linewidth=2,label='Current') #current
            ax2.set(ylabel = "Soil Moisture (% Saturation)")
            date_form = DateFormatter("%b")
            ax2.xaxis.set_major_formatter(date_form)
            
            ax2.text(0,1.01,'Soil Moisture',transform = ax2.transAxes,fontsize = 9,weight='bold') 
            ax2.text(.01,.95,'% Saturation: '+str(smcur),transform = ax2.transAxes,fontsize = 9)
        
        
        
        barmult = .05  
        #If there are reservoirs do this stuff
        if (len(res) != 0):
            if wshed!='state_of_nevada_and_eastern_sierra':
    
                idx = np.asarray([i for i in range(len(res))])
                ax4 = fig.add_subplot(224)
                #ax4.set_position([0.575,0.14, .32, .325])                        
                if len(res)>2:
                    ax4.set_position([0.59,0.145, .31, .315])
                else:
                    ax4.set_position([0.59,0.19, .31, .215])
                   
                if len(idx)==1:    
                    ax4.margins(y=1.2)
                    
                if len(idx)==2:
                    ax4.margins(y=.35)
                    b1=ax4.barh(idx-.15/2,res['res_ly_per_cap'],height = .15,color=(147/255, 147/255, 248/255),edgecolor='black',label='Previous Year',align='center')
                    b2=ax4.barh(idx+.15/2,res['res_curr_per_cap'],height = .15,color=(199/255, 250/255, 199/255),edgecolor='black',label='Current Year',align='center')
                else: 
                    b1=ax4.barh(idx-(len(idx)*barmult)/2,res['res_ly_per_cap'],height = len(idx)*barmult,color=(147/255, 147/255, 248/255),edgecolor='black',label='Previous Year',align='center')
                    b2=ax4.barh(idx+(len(idx)*barmult)/2,res['res_curr_per_cap'],height = len(idx)*barmult,color=(199/255, 250/255, 199/255),edgecolor='black',label='Current Year',align='center')
            
            
                leg=ax4.legend([b1,b2],
                            labels=['Last Year (% capacity)',
                                    'This Year (% capacity)'],
                            ncol=2,
                            prop={'size':9},
                            bbox_to_anchor=(1.16, -.08))
                
                leg.get_frame().set_linewidth(0.0)
              
                ax4.set_xlim([0,100])
                ax4.invert_yaxis()
                ax4.set_xticks([0,20,40,60,80,100])
                ax4.set_xticklabels(['0','20','40','60','80','100'])
                ax4.set_yticks(idx)
                ax4.set_yticklabels(reslabels,fontdict={'fontsize':9})
                #ax4.set_yticklabels(reslabels)
    
                ax4.text(0,1.01,'Reservoir Storage',transform = ax4.transAxes,fontsize = 9,weight='bold') 
    
                
            else:
              
                    barmult = .04
                    idx = np.asarray([i for i in range(len(res_basin_df))])
                    ax4 = fig.add_subplot(224)
                    #ax4.set_position([0.575,0.14, .32, .325])
                    ax4.set_position([0.585,0.145, .315, .315])
            
                    b1=ax4.barh(idx-(len(idx)*barmult)/2,res_basin_df['storage_basin_prev'],height= len(idx)*barmult,color=(147/255, 147/255, 248/255),edgecolor='black',label='Previous Year',align='center')
                    b2=ax4.barh(idx+(len(idx)*barmult)/2,res_basin_df['storage_basin_cur'],height= len(idx)*barmult,color=(199/255, 250/255, 199/255),edgecolor='black',label='Current Year',align='center')
                
                    leg=ax4.legend([b1,b2],
                                labels=['Last Year (% capacity)',
                                        'This Year (% capacity)'],
                                ncol=2,
                                prop={'size':9},
                                bbox_to_anchor=(1.16, -.08))
                    
                    leg.get_frame().set_linewidth(0.0)
                  
                    ax4.set_xlim([0,100])
                    ax4.invert_yaxis()
                    ax4.set_xticks([0,20,40,60,80,100])
                    ax4.set_xticklabels(['0','20','40','60','80','100'])
                    ax4.set_yticks(idx)
                    ax4.set_yticklabels(reslabels,fontdict={'size':9})
                    ax4.text(0,1.01,'Reservoir Storage',transform = ax4.transAxes,fontsize = 9,weight='bold') 
    
            #Add text to bar graph if that is unavailable or if reservoir is empty
            totals = []
            for i in ax4.patches:
                totals.append(i.get_width())
               
            for i in ax4.patches:       
                #print(i.get_y())
                if i.get_width()==.0001: #in LoadReportData I set NaN data to .0001.              
                    ax4.text(2, i.get_y()+.11,'Data Unavailable', fontsize=8,color='black')
                elif i.get_width()==0:            
                    ax4.text(2, i.get_y()+.11,'Reservoir Empty', fontsize=8,color='black')
                else:
                    ax4.text(i.get_width(), i.get_y(),'')    
    
        fig_path = fr"C:\USDA\Work\ReportDocs\Jeff_WSOR_Docs\Figs\{fname1}.png"
        print("Saving to:", fig_path)

        plt.savefig(fig_path, bbox_inches='tight',dpi=300)
            
    
        
        doc = Document()
        sections = doc.sections
        #narrow margins
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        
        
        heading = doc.add_heading('{} | {} 1, {}'.format(cleantxt,dt.date(1900,ReportMonth,1).strftime('%B'),str(ReportYear)),0)
        P = doc.add_paragraph(txt)
        #P = doc.add_paragraph()
            
        doc.add_picture(fig_path)
    
        
        title_style = heading.style
        title_style.font.name = "Arial"
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.element.xml
        title_style.font.size = Pt(16)
        rFonts = title_style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Times New Roman")
        
        
        style = doc.styles['Heading 1']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(16)
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        
        style = doc.styles['No Spacing']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        
        if len(res) !=0 or wshed=='clover_valley_and_franklin':
            p1 = doc.add_paragraph('           Statistical shading breaks at 10th, 30th, 50th, 70th, and 90th percentiles.\n           For more information visit: ')
        else:
            p1 = doc.add_paragraph('                                                                          Statistical shading breaks at 10th, 30th, 50th, 70th, and 90th percentiles.\n                                                                          For more information visit: ')
        #add a hyperlink with the normal formatting (blue underline)
        hyperlink = add_hyperlink(p1, 'https://www.nrcs.usda.gov/wps/portal/wcc/home/snowClimateMonitoring/30YearNormals/'\
                                  ,'30 year normal calculation description', '4B8BBE', True)
        
        if wshed == 'lake_tahoe' or wshed=='truckee' or wshed=='carson':
            p2 = doc.add_paragraph()
            p2 = doc.add_paragraph()
            p2.add_run('''Important Information about Forecast Coordination: ''').bold=True 
                                    
            p2.add_run('Hydrologists with the NRCS and National Weather Service California Nevada River Forecast Center (CNRFC) coordinate '\
                                   'Lake Tahoe Rise, Truckee River at Farad, Little Truckee River near Boca, and the Carson River at Ft. Churchill forecasts (following page) using output of their respective '\
                                   'hydrology models at the request of the Bureau of Reclamation. The NRCS model is a statistical model based on the current data as of the first of each month. The CNRFC '\
                                   'ensemble forecasting system incorporates near-term weather prediction and climatology into their model. These models can provide different answers because of the nature '\
                                   'of the model systems, and from the inclusion of future weather in the CNRFC model. The hydrologists agree on forecast values using guidance from both models to best '\
                                   'provide an accurate water supply forecast for these points.')
            p2.style = doc.styles['No Spacing']
            
            
        if wshed == 'spring_mountains':
            p2 = doc.add_paragraph()
            p2 = doc.add_paragraph()
            p2.add_run('SNOTEL sites in the Spring Mountains were installed in 2008. Reported percentages are based on SNOTEL medians calculated '\
                       'using data from water years 2009-2020, not the full 30-year period. Snowpack percentages in the March and April reports '\
                       'include snow course measurements from long term data collection transects.')
            p2.style = doc.styles['No Spacing']
                     
        p1.style = doc.styles['No Spacing']
    
        #doc.add_page_break()
        
        docpath = fr"C:\USDA\Work\ReportDocs\Jeff_WSOR_Docs\WordDocs\{fname1}.docx"
        pdfpath = fr"C:\USDA\Work\ReportDocs\Jeff_WSOR_Docs\PDFs\{fname1}.pdf"
        
        print('Saving '+cleantxt+' .pdf')
        print("Saving to:", docpath)
        print("Saving to:", pdfpath)

        doc.save(docpath)
        convert(docpath,pdfpath)
 
        
if __name__ == "__main__":
    main()

    
    
    
    
    
    
    
    
    
    
    
    
    

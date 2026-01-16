# -*- coding: utf-8 -*-
"""
Created on Fri Sep 10 15:08:55 2021

@author: David.Eiriksson
"""

import docx

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


from docx.shared import Pt
document = docx.Document()

from docx.shared import Pt

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(8)


p = document.add_paragraph().add_run()
#p.font.size=Pt(8)    
p = document.add_paragraph('Statistical shading breaks at 10th, 30th, 50th, 70th, and 90th percentiles.\nFor more information visit: ')

#add a hyperlink with the normal formatting (blue underline)
hyperlink = add_hyperlink(p, 'https://www.nrcs.usda.gov/wps/portal/wcc/home/snowClimateMonitoring/30YearNormals/!ut/p/z1/tVTRbpswAPwVXngEG3DA2RuNqqYhVZakTYpfKgc8YAs2NU5o9vUzi1ppSms6TfMDks2d76w7GxDwCAinx6qgqhKc7vU8JeETxhPkLRCcwwTGcDnBySacbrzFXQC2vwG3Po6neA2TmwlCMB5Povv5OoLJfQiIgf_VH-Bfwc_x4QcjNvPf_BsEiMHeajTgXwMM-vgODp3f-xzfAPiTj299_XuZoMUIBjfL6IJ_ASD_kv_MG-JvAAEk46pRJUi7LLMywRXjyoalqJkNWy46i_LcyvZVTRWzasErJWTFCxsG0DkxKi0uZE337cWCJVkrDjJjba_SZFUOUj_HY-SH1MlomDvIQztnR_V0FI3GEAcIZzQaaIW2Tcyl2_Z6A700NutqCNDfLCMg8IYkhgArNADQ6c4AKfZid34pYr4LcAGIZN-YZNI9SL1cKtW0X2xow67rXJ2wy2XWuoc2p24hjueEz9-nmjZuqep3NytFq8Dju3uAVMcVfRjXtQe2x4p14OFcC7D-yzZMYX9OcyG04-r78zOJdZv7Br_0Zv9fnZv64XXUODg5P1bXq-nPOds66ezYnbpfBE3VAw!!/dz/d5/L2dBISEvZ0FBIS9nQSEh/'\
                          ,'30 year normal calculation description', '4B8BBE', True)

    
p.style = document.styles['Normal']
#add a hyperlink with a custom color and no underline
# hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', 'FF8822', False)

document.save('demo.docx')
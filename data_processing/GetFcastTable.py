# -*- coding: utf-8 -*-
"""
Created on Fri Jan  9 13:21:16 2026

@author: Logan.Jamison
"""

def GetFcastTable(ReportMonth, ReportYear, RiseInAPI=True):
    
    import pandas as pd
    # import numpy as np
    # from datetime import datetime
    # import matplotlib.pyplot as plt
    import requests
    import os
    from datetime import datetime

    rootpath = 'C:/USDA/Work/ReportDocs/Jeff_WSOR_Docs/'
    
    domain = 'https://nwcc-apps.sc.egov.usda.gov'
    fcast_url = f'/aws-api/wsor/getFcstData?basinType=nv&pubMonth={ReportMonth}&pubYear={ReportYear}&format=json'
    fcast_json = requests.get(domain + fcast_url).json()

    rise_url = f'https://wcc.sc.egov.usda.gov/awdbRestApi/services/v1/forecasts?stationTriplets=09UTGSLR%3AUT%3ABOR%2C10337000%3ACA%3ABOR%2C%2010288500%3ANV%3ABOR&elementCodes=REST'
    rise_json = requests.get(rise_url).json()
    
    watershedlist = [  
                            'state of nevada and eastern sierra',
                            'lake tahoe',
                            'truckee',
                            'carson',
                            'walker',
                            'northern great basin',
                            'upper humboldt',
                            'lower humboldt',
                            'clover valley and franklin',
                            'snake',
                            'owyhee',
                            'eastern nevada',
                            'spring mountains',
                            'surprise valley-warner mtns',
                            #'Upper_Colorado_Region'
                      ]
        

    fcast_d = {} # Build basic dictionary with forcast data
    for watershed in watershedlist:

        #individual reservoir stats
        fcast_d[watershed] = {
                'fcst_curr': fcast_json[watershed][0]['fcst_curr'],
                'fcst_med': fcast_json[watershed][0]['fcst_med'],
                'site_meta': fcast_json[watershed][0]['site_meta']
            }
    
        # build Tahoe Rise rows and reformat to match what other fcast points will look like in table
        month_map = {
            "01": "JAN", "02": "FEB", "03": "MAR", "04": "APR",
            "05": "MAY", "06": "JUN", "07": "JUL", "08": "AUG",
            "09": "SEP", "10": "OCT", "11": "NOV", "12": "DEC"
        }
        
        
        tahoe_rows = []
        if RiseInAPI:
            for item in rise_json[1]['data']:
                if datetime.strptime(item['publicationDate'],"%Y-%m-%d %H:%M").month == ReportMonth: # only grab for current publication month
                    raw_period = item['forecastPeriod']
                    month_num = raw_period[0].split("-")[0]
                    month_abbrev = month_map[month_num]
                    level = raw_period[1]
                
                    period = f"{month_abbrev}-{level}"
                
                    values = item['forecastValues']
                
                    row = {
                        "name": rise_json[1]['forecastPointName'],
                        "Forecast Period": period,
                        "90 (KAF)": values.get("90"),
                        "70 (KAF)": values.get("70"),
                        "50 (KAF)": values.get("50"),
                        "30 (KAF)": values.get("30"),
                        "10 (KAF)": values.get("10"),
                        "30yr Median (KAF)": None,
                        "% Median": None
                    }
                
                    tahoe_rows.append(row)
                else: continue
        else:
              row = {
                  "name": "Lake Tahoe Rise Gates Closed",
                  "Forecast Period": None,
                  "90 (KAF)": None,
                  "70 (KAF)": None,
                  "50 (KAF)": None,
                  "30 (KAF)": None,
                  "10 (KAF)": None,
                  "30yr Median (KAF)": None,
                  "% Median": None
              }
          
              tahoe_rows.append(row)

    



    # function to build rows for each dataframe (individual tables)
    def build_rows(item):
        rows = []
        for period, values in item["fcst_curr"].items():
            row = {
                "name": item["name"],
                "Forecast Period": period,
            }
            for percentile, value in values.items():
                row[f"{percentile} (KAF)"] = value
            row["30yr Median (KAF)"] = item["fcst_med"].get(period)
    
            f50 = row.get("50 (KAF)") #calculate percent median from 50% value and median
            med = row.get("30yr Median (KAF)")
            if f50 is not None and med not in (None, 0):
                row["% Median"] = str(int(round(f50 / med * 100)))+"%"
            else:
                row["% Median"] = None
    
            rows.append(row)
        return rows

    dfs = {} # build dictionary of dataframes
    for watershed in watershedlist:     
        # fcast station triplets from meta data are keys
        triplet_keys = list(fcast_d[watershed]['site_meta'].keys()) #grab watershed according to its list index
    
        fcast_data = [] # build list of actual data for fcast pts in watershed
        for i in range(len(triplet_keys)): # loop through keys (triplets)
            
            fcast_data.append(
                {
                    'name': fcast_d[watershed]['site_meta'][triplet_keys[i]]['name'],
                    'fcst_curr': fcast_d[watershed]['fcst_curr'][triplet_keys[i]],
                    'fcst_med': fcast_d[watershed]['fcst_med'][triplet_keys[i]],
                   }
                )
            
        rows = [] # build individual rows using  function
        for item in fcast_data:
            rows.extend(build_rows(item))
        df = pd.DataFrame(rows) # convert to dataframe
        if watershed == 'lake tahoe': # add in GSL rise rows
            df = pd.concat([df, pd.DataFrame(tahoe_rows)], ignore_index=True)
        
        new_order = [
            "name",
            "Forecast Period",
            "90 (KAF)",
            "70 (KAF)",
            "50 (KAF)",
            "% Median",
            "30 (KAF)",
            "10 (KAF)",
            "30yr Median (KAF)"
            ]
        df = df.reindex(columns=new_order)
        df = df.rename(columns={"name": "Forecast Point","90 (KAF)": "90% (KAF)","70 (KAF)":"70% (KAF)","50 (KAF)":"50% (KAF)",
                       "30 (KAF)":"30% (KAF)","10 (KAF)":"10% (KAF)"})

        
        dfs[watershed]=df

    #%%
    
    # Create word doc table from dict of dataframes and convert to PDF
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK
    from docx2pdf import convert
    
    # asterisk = "\u20F0"
    
    # Formatting functions
    def shade_cell(cell, fill):
        """fill = hex color like 'D9D9D9' (no #)"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    
    def remove_table_borders(table): 
        tbl = table._tbl
        tblPr = tbl.tblPr
    
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'none')
            borders.append(elem)
    
        tblPr.append(borders)
        
    
    def prevent_row_split(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)
    
    # def format_vals(val): # function to remove trailing zeros from whole numbers and round unruly numbers
    #     # Leave blanks alone
    #     if val is None or val == "":
    #         return "" 
    #     try:
    #         s = str(val).strip()
    #         f = float(s)
    #         # Whole number → drop .0
    #         if f.is_integer():
    #             return str(int(f))
    #         # If there's a decimal, check how many places
    #         if "." in s:
    #             decimals = len(s.split(".")[1])
    #             # Already nicely formatted (1–2 decimals)
    #             if decimals <= 2:
    #                 return s
    #             # Too many decimals → round to 2
    #             return f"{f:.2f}"
    #         return s
    #     except Exception:
    #         return str(val)

    def format_vals(val): # function to remove trailing zeros from whole numbers
        # Leave blanks alone
        if val is None or val == "":
            return ""
        try:
            f = float(val)
            # If it's a whole number, return as int
            if f.is_integer():
                return str(int(f))
            # Otherwise keep original formatting
            return str(val)
        except:
            return str(val)
    
    def tighten_cell_paragraphs(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fmt = para.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing = 1
                    
    def set_cell_margins(table, top=0, start=50, bottom=0, end=50):
        tbl = table._tbl
        tblPr = tbl.tblPr
    
        tblCellMar = OxmlElement('w:tblCellMar')
        for m, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
    
        tblPr.append(tblCellMar)
        
    def tighten_paragraph_spacing(paragraph, before=0, after=0):
        p = paragraph.paragraph_format
        p.space_before = Pt(before)
        p.space_after = Pt(after)
    
    
    def shrink_table_font(table, size=9):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(size)
    
    # function to build word document from forecast dictionary
    def df_dict_to_word(df_dict):
        # first = True # first table
        
    
        for name, df in df_dict.items():
            
            doc = Document()
            
            # if not first: # add page break if not the first table
            #     doc.add_page_break()
            # first = False
    
            # Remove repeated forecast point names, keep only first instance
            df.loc[df["Forecast Point"].duplicated(), "Forecast Point"] = ""
    
            # Add centered title
            title_para = doc.add_paragraph()
            title_para.alignment = 1 # Centered
            run1 = title_para.add_run("Water Supply Forecast Table\n") 
            run2 = title_para.add_run("Chance that actual volume will exceed forecast") 
            run1.bold = True 
            run2.italic = True
            tighten_paragraph_spacing(title_para, before=0, after=0)
            
            # watershed name as heading
            doc.add_heading(name.title(), level=2) # title() capitalizes each word of watershed
    
            # Create table 
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"
            
            set_cell_margins(table, top=0, bottom=0, start=40, end=40)
    
            # remove_table_borders(table)
    
            # Add header cells
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                run = hdr_cells[i].paragraphs[0].add_run(str(col))
                run.bold = True
            prevent_row_split(table.rows[0])
    
            # Add data rows
            for r_idx,(_, row) in enumerate(df.iterrows()):
                row_cells = table.add_row().cells
                
                # alternating shading (skip header row)
                if r_idx % 2 == 1: 
                    for c in row_cells: shade_cell(c, "EDEDED") # light gray
                    
                for i, val in enumerate(row):
                    if pd.isna(val): #fill Nans with blanks (should only apply to lake rise)
                        val= ""
                    row_cells[i].text = format_vals(val)                   
                # add asterisk for GSL inflow only
                # if name.lower() == 'great salt lake' and row['Forecast Point'] =='Great Salt Lake Inflow':
                #    row_cells[0].text = row_cells[0].text + "*"    
                # if name.lower() == 'great salt lake' and row['Forecast Point'] =='Great Salt Lake Rise':
                #    row_cells[0].text = row_cells[0].text + "**"  
                   
                prevent_row_split(table.rows[-1])
            tighten_cell_paragraphs(table)
            #tighten_cell_paragraphs(table)
    
            # add footer to only GSL table
            # if name.lower() == 'great salt lake':
            #     foot_para = doc.add_paragraph()  
            #     run3 = foot_para.add_run('* Unlike other forecast values in this table, the GSL inflow forecast does not correct for upstream management actions, i.e. is not a "naturalized" forecast.\n')
            #     run4 = foot_para.add_run('** Units = feet. Please be advised that the lake level rise forecast for the GSL is meant to be advisory only given the uncertainty in the modeling and the preponderance of water management actions in the basin.')
    
            #     run3.font.size = Pt(10)
            #     run4.font.size = Pt(10)
               
            # tighten_paragraph_spacing(foot_para, before=0, after=2)
            
            # Add spacing between tables
            doc.add_paragraph()  
            doc.add_paragraph()
            
            
            shrink_table_font(table, size=9)
            
            output_path = rootpath + f"WordDocs/{name}_forecast_table.docx"
            # Save the Word document
            doc.save(output_path)
            
            # convert to pdf
            fcast_pdf = rootpath + f"PDFs/{name}_forecast_table.pdf"
            convert(output_path,fcast_pdf)
    
    # output_docx = rootpath + "WordDocs/" + "All_Forecast_Tables.docx"
    
    df_dict_to_word(dfs)
    
    # fcast_pdf = rootpath +"PDFs/96A_ForecastSummaryTable.pdf"
    # convert(output_docx,fcast_pdf)
    
    
    
    
    
              
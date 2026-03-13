# -*- coding: utf-8 -*-
"""
Created on Fri Mar 13 13:27:20 2026

@author: Logan.Jamison
"""

def StatewideResSum_WSOR_NV(wy,reportmonth,prec_df_api,basin_res_df_api,res_d):
    from LoadReportData_NV import LoadReportData_NV
    import pandas as pd
    # wy = 2025
    # reportmonth = 2
    
    [_,_,_,res,_,_,_,title,_,_,_]=LoadReportData_NV('state_of_nevada_and_eastern_sierra',wy,reportmonth,prec_df_api,basin_res_df_api,res_d)
         
    
    # Make Dataframe/table for all reservoirs from res dataframe
   
    # Check for missing data by trying to convert data to integers
    for col in res.columns:
        # Try converting the column to numeric (temporarily)
        converted = pd.to_numeric(res[col], errors='coerce')

        # Find rows that failed conversion
        bad_rows = converted[converted.isna()].index

        if len(bad_rows) > 0:
            raise ValueError(
                "You are missing reservoir data!\n"
                f"Missing values found for: {bad_rows.tolist()}"
            )

        # If no errors, assign the converted int column back
        #res[col] = converted.astype(int)
        
    res[['res_curr_per_cap','res_ly_per_cap']] = res[['res_curr_per_cap','res_ly_per_cap']].astype(int)    
    
    cols = ['res_cap', 'res_curr']

    for c in cols:
        res[c] = res[c].apply(lambda x: round(x,1) if x < 10 else int(x))

        
    res['Reservoir'] = res.index
    
    res = res.rename(columns={'res_curr': 'Current Storage (KAF)','res_cap': 'Reservoir Capacity (KAF)','res_ly_per_cap': 'Last Yr % Capacity', 'res_curr_per_cap': 'This Yr % Capacity'})   
    res = res[['Reservoir','Current Storage (KAF)','Reservoir Capacity (KAF)','Last Yr % Capacity','This Yr % Capacity']]
    res = res.sort_index()
          
    #Tidy up Smith and Morehouse Reservoir, so it's not on two lines.
    res.at['Smith And Morehouse Reservoir','Reservoir']='Smith and Morehouse'

    # Keep formatting of numerical values in res tables
    
    def format_value(res):
        # Keep integers as integers
        if isinstance(res, int):
            return str(res)
        # Floats: remove trailing .0
        if isinstance(res, float):
            if res.is_integer():
                return str(int(res))
            else:
                return f"{res:.1f}"   # or whatever precision you want
        return str(res)
    
    res = res.applymap(format_value)
    
        
        #%%
    import docx
    from docx.shared import Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx2pdf import convert
    from docx.shared import RGBColor
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    rootpath = 'C:/USDA/Work/ReportDocs/Jeff_WSOR_Docs/'
    doc = docx.Document()
    
    sections = doc.sections
    #narrow margins
    for section in sections:
        section.top_margin = Cm(.25)
        section.bottom_margin = Cm(.25)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1)
    
    
    if reportmonth == 1:
        MonthText = "January 1, " + str(wy)
    if reportmonth == 2:
        MonthText = "Feb 1, " + str(wy)
    if reportmonth == 3:
        MonthText = "March 1, " + str(wy)
    if reportmonth == 4:
        MonthText = "April 1, " + str(wy)
    if reportmonth == 5:
        MonthText = "May 1, " + str(wy)
    if reportmonth == 6:
        MonthText = "June 1, " + str(wy)
    if reportmonth == 7:
        MonthText = "July 1, " + str(wy)
    if reportmonth == 8:
        MonthText = "Aug 1, " + str(wy)
    if reportmonth == 9:
        MonthText = "Sept 1, " + str(wy)
    if reportmonth == 10:
        MonthText = "Oct 1, " + str(wy)
    if reportmonth == 11:
        MonthText = "Nov 1, " + str(wy-1)
    if reportmonth == 12:
        MonthText = "Dec 1, " + str(wy-1)
      
    
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    
    doc.add_heading(MonthText, 1)
    
    
    # t = doc.add_table(res_basin_df.shape[0]+1, res_basin_df.shape[1])
    # t.style = 'Medium Shading 1'
    # #'Light Shading Accent 1'
    # #'Medium Shading 1 Accent 1'
    # # add the header rows.
    
    
    # for j in range(res_basin_df.shape[-1]):
    #     t.cell(0,j).text = res_basin_df.columns[j]
    
    # # add the rest of the data frame
    # for i in range(res_basin_df.shape[0]):
    #     for j in range(res_basin_df.shape[-1]):
    #         t.cell(i+1,j).text = str(res_basin_df.values[i,j])
    
    
    # for row in t.rows:
    #     row.height = Cm(0.31)
    
    # for row in t.rows:
    #     for cell in row.cells:
    #         paragraphs = cell.paragraphs
    #         for paragraph in paragraphs:
    #             for run in paragraph.runs:
    #                 font = run.font
    #                 font.size= Pt(7.25)
    #                 font.name = 'Arial'
    
    
    # pct_change = res_basin_df['Last Yr % Capacity \n (Basinwide)']-res_basin_df['This Yr % Capacity \n (Basinwide)']
    
    # for i in range(len(res_basin_df)):
    #     if pct_change[i]>5:
    #         shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D0312D"/>'.format(nsdecls('w')))
    #         t.rows[i+1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)    
    #     elif pct_change[i]<-5:
    #         shading_elm_1 = parse_xml(r'<w:shd {} w:fill="3CB043"/>'.format(nsdecls('w')))
    #         t.rows[i+1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1) 
    
    # Footer = 'Red (green) shading indicates >5% decrease (increase) in % capacity from this time last year.'  
    Footer = ''
    P = doc.add_paragraph().add_run(Footer)    
    P.font.size = Pt(8)
    P.italic = True
    P.font.name = "Arial"
    
    
    #doc.add_paragraph()
    
    t = doc.add_table(res.shape[0]+1, res.shape[1])
    t.style = 'Medium Shading 1'
    #'Light Shading Accent 1'
    #'Medium Shading 1 Accent 1'
    # add the header rows.
    for j in range(res.shape[-1]):
        t.cell(0,j).text = res.columns[j]
        
    # add the rest of the data frame
    for i in range(res.shape[0]):
        for j in range(res.shape[-1]):
            t.cell(i+1,j).text = str(res.values[i,j])
    
    for row in t.rows:
        row.height = Cm(0.31)
        row.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(7.25)
                    font.name = 'Arial'
                    
    # pct_change = res['Last Yr % Capacity']-res['This Yr % Capacity']
    
    # for i in range(len(res)):
    #     if pct_change[i]>5:
    #         shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D0312D"/>'.format(nsdecls('w')))
    #         t.rows[i+1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)    
    #     elif pct_change[i]<-5:
    #         shading_elm_1 = parse_xml(r'<w:shd {} w:fill="3CB043"/>'.format(nsdecls('w')))
    #         t.rows[i+1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_1)
    
    
    
    
    
    #Footer = 'Red (green) shading indicates >5% decrease (increase) in % capacity from this time last year.'  
    Footer = ''
    P = doc.add_paragraph().add_run(Footer)    
    P.font.size = Pt(8)
    P.italic = True
    P.font.name = "Arial"
    
    
    
    
    # WAIText = [] 
    
    
    # P = doc.add_paragraph().add_run(WAIText)
    # P.font.size = Pt(9)
    # P.font.name = 'Arial'
    
    # save the doc
    WAIdoc = rootpath+'WordDocs/96_Res_Summary.docx'
    WAIpdf = rootpath+'PDFs/96_Res_Summary.pdf'
    doc.save(WAIdoc)
    
    print('Generating Reservoir Summary PDF.  This takes a minute.  Be patient.')
    convert(WAIdoc,WAIpdf)